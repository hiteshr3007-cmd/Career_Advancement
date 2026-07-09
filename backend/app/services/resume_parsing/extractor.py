import io

import pdfplumber
from docx import Document


class UnsupportedResumeFormat(Exception):
    pass


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """Extract raw text from a resume file. Supports pdf, docx. Legacy .doc has
    best-effort support only (recommend the client convert to pdf/docx)."""
    file_type = file_type.lower().lstrip(".")

    if file_type == "pdf":
        return _extract_pdf(file_bytes)
    if file_type == "docx":
        return _extract_docx(file_bytes)
    if file_type == "doc":
        return _extract_legacy_doc(file_bytes)

    raise UnsupportedResumeFormat(f"Unsupported resume file type: {file_type}")


def _extract_pdf(file_bytes: bytes) -> str:
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def _extract_docx(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


def _extract_legacy_doc(file_bytes: bytes) -> str:
    # No pure-python .doc (binary OLE) parser is bundled. Best effort: strip
    # non-printable bytes so downstream LLM parsing still has a chance.
    text = file_bytes.decode("latin-1", errors="ignore")
    printable = "".join(ch if ch.isprintable() or ch in "\n\r\t" else " " for ch in text)
    collapsed = " ".join(printable.split())
    return collapsed
